"""Screen Tools — screenshot capture, OCR (Apple Vision), vision model queries.

On-command only. FRIDAY never watches the screen unless the user explicitly asks.
Uses macOS Vision framework for OCR (free, fast, offline) and Qwen2.5-VL via
Ollama for image understanding (what's on screen, explain UI, read errors).
"""

import asyncio
import base64
import os
import re
import subprocess
from datetime import datetime
from pathlib import Path

from friday.core.types import ToolResult, ToolError, ErrorCode, Severity

FRIDAY_SCREENSHOTS = Path.home() / "Downloads" / "friday_screenshots"
_SCREENSHOT_TTL_HOURS = 48  # Auto-delete screenshots older than 2 days


def _screen_access_enabled() -> bool:
    return os.environ.get("FRIDAY_SCREEN_ACCESS", "").lower() == "true"


def _cleanup_old_screenshots():
    """Delete screenshots older than _SCREENSHOT_TTL_HOURS. Runs on every capture."""
    if not FRIDAY_SCREENSHOTS.exists():
        return
    import time
    cutoff = time.time() - (_SCREENSHOT_TTL_HOURS * 3600)
    for f in FRIDAY_SCREENSHOTS.glob("screen_*.png"):
        try:
            if f.stat().st_mtime < cutoff:
                f.unlink()
        except OSError:
            pass


async def capture_screen(
    region: dict = None,
) -> ToolResult:
    """Take a screenshot. Optionally capture a region: {"x": 0, "y": 0, "w": 800, "h": 600}.

    Returns file path + base64-encoded image for vision model.
    """
    if not _screen_access_enabled():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.PERMISSION_DENIED,
                message="Screen access disabled. Set FRIDAY_SCREEN_ACCESS=true in .env to enable.",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    FRIDAY_SCREENSHOTS.mkdir(parents=True, exist_ok=True)
    _cleanup_old_screenshots()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = FRIDAY_SCREENSHOTS / f"screen_{ts}.png"

    try:
        if region:
            cmd = [
                "screencapture", "-x",
                "-R", f"{region['x']},{region['y']},{region['w']},{region['h']}",
                str(path),
            ]
        else:
            cmd = ["screencapture", "-x", str(path)]

        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await asyncio.wait_for(proc.communicate(), timeout=10)

        if proc.returncode != 0:
            return ToolResult(
                success=False,
                error=ToolError(
                    code=ErrorCode.COMMAND_FAILED,
                    message=f"Screenshot failed: {stderr.decode(errors='replace')}",
                    severity=Severity.MEDIUM,
                    recoverable=True,
                ),
            )

        image_bytes = path.read_bytes()
        image_b64 = base64.b64encode(image_bytes).decode("utf-8")

        return ToolResult(
            success=True,
            data={
                "saved_path": str(path),
                "image_b64": image_b64,
                "media_type": "image/png",
                "size_bytes": len(image_bytes),
                "timestamp": ts,
            },
        )
    except asyncio.TimeoutError:
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.TIMEOUT,
                message="Screenshot timed out.",
                severity=Severity.MEDIUM,
                recoverable=True,
            ),
        )
    except Exception as e:
        return ToolResult(success=False, error=ToolError(
            code=ErrorCode.COMMAND_FAILED, message=str(e),
            severity=Severity.MEDIUM, recoverable=True,
        ))


async def ocr_screen(image_path: str = None) -> ToolResult:
    """Extract all text from a screenshot using Apple Vision framework (offline, fast).

    If no image_path provided, takes a fresh screenshot first.
    """
    if not _screen_access_enabled():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.PERMISSION_DENIED,
                message="Screen access disabled. Set FRIDAY_SCREEN_ACCESS=true in .env to enable.",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    # Take a fresh screenshot if no path given
    if not image_path:
        screenshot = await capture_screen()
        if not screenshot.success:
            return screenshot
        image_path = screenshot.data["saved_path"]

    if not Path(image_path).exists():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.FILE_NOT_FOUND,
                message=f"Image not found: {image_path}",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    # Apple Vision OCR via Swift (fast, accurate, offline)
    swift_code = f'''
import Vision
import AppKit

let path = "{image_path}"
guard let image = NSImage(contentsOfFile: path) else {{ fputs("ERROR: Could not load image\\n", stderr); exit(1) }}
guard let tiffData = image.tiffRepresentation else {{ fputs("ERROR: No TIFF representation\\n", stderr); exit(1) }}

let handler = VNImageRequestHandler(data: tiffData, options: [:])
let request = VNRecognizeTextRequest()
request.recognitionLevel = .accurate
request.usesLanguageCorrection = true

do {{
    try handler.perform([request])
}} catch {{
    fputs("ERROR: Vision request failed: \\(error)\\n", stderr)
    exit(1)
}}

guard let results = request.results else {{ exit(0) }}
for observation in results {{
    if let candidate = observation.topCandidates(1).first {{
        print(candidate.string)
    }}
}}
'''

    try:
        proc = await asyncio.create_subprocess_exec(
            "swift", "-e", swift_code,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)

        if proc.returncode != 0:
            err_text = stderr.decode("utf-8", errors="replace").strip()
            return ToolResult(
                success=False,
                error=ToolError(
                    code=ErrorCode.COMMAND_FAILED,
                    message=f"OCR failed: {err_text}",
                    severity=Severity.MEDIUM,
                    recoverable=True,
                ),
            )

        text = stdout.decode("utf-8", errors="replace").strip()

        return ToolResult(
            success=True,
            data={
                "text": text,
                "image_path": image_path,
                "char_count": len(text),
                "line_count": len(text.splitlines()) if text else 0,
            },
        )
    except asyncio.TimeoutError:
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.TIMEOUT,
                message="OCR timed out after 30s.",
                severity=Severity.MEDIUM,
                recoverable=True,
            ),
        )
    except Exception as e:
        return ToolResult(success=False, error=ToolError(
            code=ErrorCode.COMMAND_FAILED, message=str(e),
            severity=Severity.MEDIUM, recoverable=True,
        ))


async def ask_about_screen(
    query: str,
    image_path: str = None,
) -> ToolResult:
    """Look at the screen and answer a question about what's visible.

    Uses Qwen2.5-VL via Ollama for image understanding. If no image_path,
    takes a fresh screenshot. Falls back to OCR if vision model unavailable.
    """
    if not _screen_access_enabled():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.PERMISSION_DENIED,
                message="Screen access disabled. Set FRIDAY_SCREEN_ACCESS=true in .env to enable.",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    # Take a fresh screenshot if no path given
    if not image_path:
        screenshot = await capture_screen()
        if not screenshot.success:
            return screenshot
        image_path = screenshot.data["saved_path"]
        image_b64 = screenshot.data["image_b64"]
    else:
        if not Path(image_path).exists():
            return ToolResult(
                success=False,
                error=ToolError(
                    code=ErrorCode.FILE_NOT_FOUND,
                    message=f"Image not found: {image_path}",
                    severity=Severity.MEDIUM,
                    recoverable=False,
                ),
            )
        image_b64 = base64.b64encode(Path(image_path).read_bytes()).decode("utf-8")

    prompt = query or (
        "Describe what's on this screen. What application is open? "
        "What content is visible? Any important information?"
    )

    # Try Qwen2.5-VL via Ollama
    try:
        import ollama

        response = ollama.chat(
            model="qwen2.5vl:7b",
            messages=[{
                "role": "user",
                "content": prompt,
                "images": [image_b64],
            }],
        )

        return ToolResult(
            success=True,
            data={
                "answer": response["message"]["content"],
                "image_path": image_path,
                "model": "qwen2.5vl:7b",
            },
        )

    except Exception as vision_err:
        # Fallback: OCR the screen + use text LLM to answer
        ocr_result = await ocr_screen(image_path=image_path)
        if not ocr_result.success:
            return ToolResult(
                success=False,
                error=ToolError(
                    code=ErrorCode.COMMAND_FAILED,
                    message=f"Vision model unavailable ({vision_err}) and OCR also failed.",
                    severity=Severity.MEDIUM,
                    recoverable=True,
                ),
            )

        screen_text = ocr_result.data["text"]
        if not screen_text:
            return ToolResult(
                success=True,
                data={
                    "answer": "I can see the screen but there's no readable text on it. "
                              "The vision model (qwen2.5vl) isn't available — pull it with "
                              "`ollama pull qwen2.5vl:7b` for full image understanding.",
                    "image_path": image_path,
                    "model": "ocr_fallback",
                },
            )

        # Use text LLM to answer from OCR text
        try:
            from friday.core.llm import cloud_chat, extract_text

            messages = [
                {"role": "system", "content": (
                    "You are FRIDAY. The user asked you to look at their screen. "
                    "Below is all the text extracted from the screenshot via OCR. "
                    "Answer their question based on what you can see in the text. "
                    "Be direct and specific."
                )},
                {"role": "user", "content": (
                    f"Screen text:\n---\n{screen_text[:4000]}\n---\n\n"
                    f"The user's question: {prompt}"
                )},
            ]

            resp = cloud_chat(messages=messages, max_tokens=400)
            answer = extract_text(resp)

            return ToolResult(
                success=True,
                data={
                    "answer": answer,
                    "image_path": image_path,
                    "model": "ocr_fallback",
                    "ocr_text": screen_text[:2000],
                },
            )
        except Exception:
            # Last resort: just return the OCR text
            return ToolResult(
                success=True,
                data={
                    "answer": f"Here's what I can read on screen:\n\n{screen_text[:3000]}",
                    "image_path": image_path,
                    "model": "ocr_raw",
                },
            )


async def _ocr_image(image_path: str) -> str:
    """Run Apple Vision OCR on an image, return raw text. Internal helper."""
    swift_code = f'''
import Vision
import AppKit

let path = "{image_path}"
guard let image = NSImage(contentsOfFile: path) else {{ fputs("ERROR: Could not load image\\n", stderr); exit(1) }}
guard let tiffData = image.tiffRepresentation else {{ fputs("ERROR: No TIFF representation\\n", stderr); exit(1) }}

let handler = VNImageRequestHandler(data: tiffData, options: [:])
let request = VNRecognizeTextRequest()
request.recognitionLevel = .accurate
request.usesLanguageCorrection = true

do {{
    try handler.perform([request])
}} catch {{
    fputs("ERROR: Vision request failed: \\(error)\\n", stderr)
    exit(1)
}}

guard let results = request.results else {{ exit(0) }}
for observation in results {{
    if let candidate = observation.topCandidates(1).first {{
        print(candidate.string)
    }}
}}
'''
    proc = await asyncio.create_subprocess_exec(
        "swift", "-e", swift_code,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, _ = await asyncio.wait_for(proc.communicate(), timeout=30)
    return stdout.decode("utf-8", errors="replace").strip()


async def _activate_app(app_name: str):
    """Bring an app to the foreground via AppleScript."""
    script = f'tell application "{app_name}" to activate'
    proc = await asyncio.create_subprocess_exec(
        "osascript", "-e", script,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    await proc.communicate()
    await asyncio.sleep(1.0)  # Let the app come to front and render


async def _get_frontmost_window_bounds() -> tuple[int, int, int, int] | None:
    """Get the frontmost app's main window (x, y, width, height).

    Picks the largest window by area — avoids grabbing tiny toolbar/tab-bar
    windows that some apps (Safari, etc.) report as window 1.
    """
    script = '''
tell application "System Events"
    set fp to first process whose frontmost is true
    set bestArea to 0
    set bestResult to ""
    repeat with w in windows of fp
        set {x, y} to position of w
        set {wi, hi} to size of w
        set a to wi * hi
        if a > bestArea then
            set bestArea to a
            set bestResult to (x as text) & "," & (y as text) & "," & (wi as text) & "," & (hi as text)
        end if
    end repeat
    return bestResult
end tell
'''
    proc = await asyncio.create_subprocess_exec(
        "osascript", "-e", script,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, _ = await proc.communicate()
    try:
        parts = stdout.decode().strip().split(",")
        return int(parts[0]), int(parts[1]), int(parts[2]), int(parts[3])
    except (ValueError, IndexError):
        return None


async def _capture_window(save_path: str) -> bool:
    """Capture just the frontmost window (not full screen). Falls back to full screen."""
    bounds = await _get_frontmost_window_bounds()

    if bounds:
        x, y, w, h = bounds
        # Offset y slightly to skip the title bar area, capture content area
        # macOS title bar is ~28px but we keep it to avoid cutting content
        cmd = ["screencapture", "-x", "-R", f"{x},{y},{w},{h}", save_path]
    else:
        cmd = ["screencapture", "-x", save_path]

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    await asyncio.wait_for(proc.communicate(), timeout=10)
    return Path(save_path).exists()


async def _scroll_to_top():
    """Scroll to the very top of the page. Cmd+Up (works in browsers, docs, PDFs)."""
    script = '''
tell application "System Events"
    key code 126 using {command down}
end tell
'''
    proc = await asyncio.create_subprocess_exec(
        "osascript", "-e", script,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    await proc.communicate()
    await asyncio.sleep(0.8)


async def _scroll_page_down():
    """Scroll down in the frontmost app. Uses arrow down x15 (more reliable than Page Down)."""
    # Arrow down keystrokes work in virtually every app
    # 15 arrow-downs ≈ 1 page of content
    script = '''
tell application "System Events"
    repeat 15 times
        key code 125
    end repeat
end tell
'''
    proc = await asyncio.create_subprocess_exec(
        "osascript", "-e", script,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    await proc.communicate()
    await asyncio.sleep(0.8)  # Let the app render after scroll


def _filter_content_lines(text: str) -> list[str]:
    """Filter OCR text to only content-relevant lines (skip UI chrome).

    Removes short UI fragments, browser toolbars, OS chrome, etc.
    Used for overlap comparison so UI elements don't inflate match ratio.
    """
    lines = text.strip().splitlines()
    filtered = []
    for line in lines:
        stripped = line.strip()
        # Skip very short lines (browser buttons, icons, single chars)
        if len(stripped) < 5:
            continue
        # Skip common browser/OS UI patterns
        if re.match(r'^(File|Edit|View|Insert|Format|Tools|Window|Help|Extensions?)\s', stripped):
            continue
        if re.match(r'^https?://', stripped):
            continue
        # Skip lines that look like tab titles with dots/ellipsis
        if stripped.count('...') >= 2 or stripped.count('…') >= 2:
            continue
        filtered.append(stripped)
    return filtered


def _text_overlap(prev: str, curr: str) -> float:
    """Estimate how much of curr overlaps with prev (0.0 to 1.0).

    Filters UI chrome lines first so only content is compared.
    """
    if not prev or not curr:
        return 0.0
    prev_lines = set(_filter_content_lines(prev)[-20:])
    curr_lines = _filter_content_lines(curr)[:20]
    if not prev_lines or not curr_lines:
        return 0.0
    overlap = sum(1 for line in curr_lines if line in prev_lines)
    return overlap / len(curr_lines) if curr_lines else 0.0


def _dedupe_text(prev: str, curr: str) -> str:
    """Remove overlapping lines from curr that already appeared at the end of prev."""
    prev_lines = prev.strip().splitlines()
    curr_lines = curr.strip().splitlines()
    if not prev_lines or not curr_lines:
        return curr

    # Build set of last 30 lines from prev for matching
    tail_set = set(prev_lines[-30:])

    # Find the longest prefix of curr that overlaps with prev's tail
    # Walk through curr_lines until we find a line NOT in the tail
    best_skip = 0
    for i, cline in enumerate(curr_lines):
        stripped = cline.strip()
        if not stripped or len(stripped) < 5:
            # Short/empty lines — skip them in the overlap zone
            if i == best_skip:
                best_skip = i + 1
            continue
        if stripped in tail_set or cline in tail_set:
            best_skip = i + 1
        else:
            # First non-overlapping content line — stop
            break

    return "\n".join(curr_lines[best_skip:])


async def capture_full_page(max_scrolls: int = 20, app: str = None) -> ToolResult:
    """Capture the FULL page by scrolling and OCR-ing each viewport.

    Scrolls the frontmost app page-by-page, OCRs each viewport,
    deduplicates overlapping text, and returns the complete content.
    Works with any scrollable app — browsers, PDFs, Word docs, etc.

    Args:
        max_scrolls: Maximum page-downs to attempt.
        app: Optional app name to activate first (e.g. "Safari", "Preview").
    """
    if not _screen_access_enabled():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.PERMISSION_DENIED,
                message="Screen access disabled. Set FRIDAY_SCREEN_ACCESS=true in .env to enable.",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    # Activate the target app if specified
    if app:
        await _activate_app(app)

    # Click center of the frontmost window to ensure content area has focus
    bounds = await _get_frontmost_window_bounds()
    if bounds:
        cx, cy = bounds[0] + bounds[2] // 2, bounds[1] + bounds[3] // 2
        click_script = f'''
tell application "System Events"
    click at {{{cx}, {cy}}}
end tell
'''
        proc = await asyncio.create_subprocess_exec(
            "osascript", "-e", click_script,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        await proc.communicate()
        await asyncio.sleep(0.3)

    # Scroll to top first so we capture everything from the beginning
    await _scroll_to_top()

    FRIDAY_SCREENSHOTS.mkdir(parents=True, exist_ok=True)
    _cleanup_old_screenshots()

    all_text_parts = []
    prev_text = ""
    pages_captured = 0
    consecutive_high_overlap = 0  # Need 2 high-overlap frames to confirm end

    for i in range(max_scrolls):
        # Capture just the frontmost window (not full screen)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S") + f"_{i}"
        path = FRIDAY_SCREENSHOTS / f"screen_{ts}.png"

        captured = await _capture_window(str(path))
        if not captured:
            break

        # OCR this viewport
        text = await _ocr_image(str(path))

        # Clean up the screenshot immediately
        try:
            path.unlink()
        except OSError:
            pass

        if not text.strip():
            if i == 0:
                break
            break

        # Check if we've stopped scrolling (same content as last page)
        if i > 0:
            overlap = _text_overlap(prev_text, text)
            if overlap > 0.85:
                consecutive_high_overlap += 1
                if consecutive_high_overlap >= 2:
                    # Confirmed bottom — two frames in a row with >85% overlap
                    break
            else:
                consecutive_high_overlap = 0

        # Deduplicate and append
        if i == 0:
            all_text_parts.append(text)
        else:
            deduped = _dedupe_text(prev_text, text)
            if deduped.strip():
                all_text_parts.append(deduped)

        prev_text = text
        pages_captured += 1

        # Scroll down for next page
        await _scroll_page_down()

    full_text = "\n".join(all_text_parts)

    return ToolResult(
        success=True,
        data={
            "text": full_text,
            "pages_captured": pages_captured,
            "char_count": len(full_text),
            "line_count": len(full_text.splitlines()) if full_text else 0,
        },
    )


async def read_screen(
    app: str = None,
    full_page: bool = False,
) -> ToolResult:
    """Read what's on screen and return clean, structured text content.

    General-purpose screen reader — not just for questions. Use this to:
    - Read a job posting visible on screen
    - Read an article or document
    - Read form content or any page
    - Understand what app is showing before taking action

    Args:
        app: App to activate first (e.g. "Safari", "Preview").
        full_page: If True, scrolls through the entire page. If False, reads current viewport only.
    """
    if not _screen_access_enabled():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.PERMISSION_DENIED,
                message="Screen access disabled. Set FRIDAY_SCREEN_ACCESS=true in .env to enable.",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    if full_page:
        result = await capture_full_page(app=app)
        if not result.success:
            return result
        raw_text = result.data["text"]
        pages = result.data["pages_captured"]
    else:
        if app:
            await _activate_app(app)
        ocr_result = await ocr_screen()
        if not ocr_result.success:
            return ocr_result
        raw_text = ocr_result.data["text"]
        pages = 1

    # Clean the text — remove UI chrome
    clean_lines = []
    for line in raw_text.splitlines():
        stripped = line.strip()
        if len(stripped) < 3:
            continue
        if re.match(r'^(File|Edit|View|Insert|Format|Tools|Window|Help|Extensions?)\s', stripped):
            continue
        if re.match(r'^https?://', stripped):
            continue
        if stripped.count('...') >= 2 or stripped.count('…') >= 2:
            continue
        clean_lines.append(line)
    clean_text = "\n".join(clean_lines)

    return ToolResult(
        success=True,
        data={
            "text": clean_text,
            "pages_captured": pages,
            "char_count": len(clean_text),
            "line_count": len(clean_text.splitlines()) if clean_text else 0,
        },
    )


def _markdown_to_docx(doc, text: str):
    """Parse markdown text and add properly formatted paragraphs to a docx document.

    Handles: headings (#/##/###), bold (**text**), italic (*text*),
    numbered lists, bullet lists, inline bold within paragraphs.
    """
    from docx.shared import Pt

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=2)
            i += 1
            continue

        # Horizontal rule (---) — skip
        stripped = line.strip()
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Bold-only line (acts as subheading) — strip ** from display
        if stripped.startswith("**") and stripped.endswith("**"):
            clean = stripped.strip("* ")
            h = doc.add_heading(clean, level=3)
            h.style.font.size = Pt(12)
            i += 1
            continue

        # Bullet list item
        if re.match(r'^[\-\*]\s', stripped):
            para = doc.add_paragraph(style="List Bullet")
            _add_rich_text(para, stripped[2:].strip())
            i += 1
            continue

        # Numbered list item (e.g. "1. ", "1) ", "a. ", "a) ")
        num_match = re.match(r'^(\d+[\.\)]|[a-z][\.\)])\s', stripped)
        if num_match:
            para = doc.add_paragraph(style="List Number")
            _add_rich_text(para, stripped[num_match.end():].strip())
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip()
            if not next_line.strip():
                break
            if next_line.startswith("#") or re.match(r'^[\-\*]\s', next_line.strip()):
                break
            if re.match(r'^(\d+[\.\)]|[a-z][\.\)])\s', next_line.strip()):
                break
            para_lines.append(next_line)
            i += 1

        full_text = " ".join(l.strip() for l in para_lines)
        para = doc.add_paragraph()
        _add_rich_text(para, full_text)


def _add_rich_text(paragraph, text: str):
    """Add text with inline bold (**text**) and italic (*text*) to a paragraph."""
    # Split on bold markers first
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italic within non-bold parts
            sub_parts = re.split(r'(\*[^*]+\*)', part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    if sp:
                        paragraph.add_run(sp)


async def solve_screen_questions(
    save_path: str = None,
    app: str = None,
    full_page: bool = True,
) -> ToolResult:
    """Capture the screen, read all questions, solve them, save answers to .docx.

    Args:
        save_path: Where to save the .docx. Defaults to ~/Documents/friday_files/
        app: App name to activate first (e.g. "Safari", "Preview", "Google Chrome").
        full_page: If True (default), scrolls from top to bottom capturing everything.
                   If False, only captures the current viewport — use when the user
                   says "just this screen" or "solve what's on my screen right now".
    """
    if not _screen_access_enabled():
        return ToolResult(
            success=False,
            error=ToolError(
                code=ErrorCode.PERMISSION_DENIED,
                message="Screen access disabled. Set FRIDAY_SCREEN_ACCESS=true in .env to enable.",
                severity=Severity.MEDIUM,
                recoverable=False,
            ),
        )

    # Step 1: Capture text — full page (scrolling) or current viewport only
    if full_page:
        capture_result = await capture_full_page(app=app)
        if not capture_result.success:
            return capture_result
        screen_text = capture_result.data["text"]
        pages = capture_result.data["pages_captured"]
    else:
        # Viewport-only: activate app, screenshot, OCR — no scrolling
        if app:
            await _activate_app(app)
        ocr_result = await ocr_screen()
        if not ocr_result.success:
            return ocr_result
        screen_text = ocr_result.data["text"]
        pages = 1

    if not screen_text.strip():
        return ToolResult(
            success=True,
            data={
                "answer": "I captured the screen but couldn't read any text. "
                          "Make sure the content is visible and try again.",
            },
        )

    # Step 2: Clean OCR text — strip obvious UI chrome before sending to LLM
    clean_lines = []
    for line in screen_text.splitlines():
        stripped = line.strip()
        # Skip very short UI fragments
        if len(stripped) < 3:
            continue
        # Skip browser toolbar / menu bar patterns
        if re.match(r'^(File|Edit|View|Insert|Format|Tools|Window|Help|Extensions?)\s', stripped):
            continue
        if re.match(r'^https?://', stripped):
            continue
        if stripped.count('...') >= 2 or stripped.count('…') >= 2:
            continue
        clean_lines.append(line)
    clean_text = "\n".join(clean_lines)

    # Step 3: Send to LLM to solve
    from friday.core.llm import cloud_chat, extract_text

    # Keep input under ~4K tokens so Groq free tier doesn't reject it
    # If too long, take first 5K + last 3K chars to catch questions at start and end
    if len(clean_text) > 8000:
        clean_text = clean_text[:5000] + "\n\n[...middle content...]\n\n" + clean_text[-3000:]

    messages = [
        {"role": "system", "content": (
            "You are an expert tutor and problem solver. "
            "Below is text from the user's screen showing questions or problems.\n\n"
            "Your job:\n"
            "1. Identify every question, problem, or exercise\n"
            "2. Solve each one completely with clear working/explanation\n"
            "3. Number your answers to match the question numbers if visible\n"
            "4. For multiple choice: state the correct answer AND explain why\n"
            "5. For calculations: show your working step by step\n"
            "6. For essays/written: provide a complete, well-structured answer\n\n"
            "FORMATTING:\n"
            "- Use ## for major question group headings\n"
            "- Use **Question X.X** as bold subheading for each answer\n"
            "- Use numbered lists (1. 2. 3.) for step-by-step solutions\n"
            "- Use **bold** for final answers and key terms\n"
            "- Start directly with the first question — no preamble\n"
            "- Give THOROUGH, DETAILED answers. Each answer should be a full paragraph minimum.\n\n"
            "Be thorough. Answer EVERY question you find."
        )},
        {"role": "user", "content": (
            f"Here is the page content ({pages} page{'s' if pages != 1 else ''} captured):\n\n"
            f"---\n{clean_text}\n---\n\n"
            "Find and solve every question on this page. Give detailed answers."
        )},
    ]

    resp = cloud_chat(messages=messages, max_tokens=8000)
    answers = extract_text(resp)

    # Step 3: Save to .docx with proper formatting
    if not save_path:
        save_dir = Path.home() / "Documents" / "friday_files"
        save_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_path = str(save_dir / f"Screen_Answers_{ts}.docx")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"
        style.paragraph_format.space_after = Pt(6)

        # Title
        title_para = doc.add_heading("Screen Questions — Solved by FRIDAY", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        now_str = datetime.now().strftime("%B %d, %Y at %H:%M")
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f"Captured {pages} page{'s' if pages != 1 else ''} — {now_str}")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # Spacer

        # Answers section — properly formatted from markdown
        _markdown_to_docx(doc, answers)

        doc.save(save_path)

    except ImportError:
        # Fallback: save as .txt
        save_path = save_path.replace(".docx", ".txt")
        Path(save_path).write_text(
            f"Screen Questions — Solved by FRIDAY\n"
            f"{'=' * 40}\n\n"
            f"ANSWERS:\n\n{answers}",
            encoding="utf-8",
        )

    return ToolResult(
        success=True,
        data={
            "answers": answers,
            "save_path": save_path,
            "pages_captured": pages,
            "questions_text_length": len(screen_text),
            "answers_length": len(answers),
        },
    )


# ── Tool schemas for dispatch ──────────────────────────────────────────────

TOOL_SCHEMAS = {
    "capture_screen": {
        "fn": capture_screen,
        "schema": {
            "name": "capture_screen",
            "description": "Take a screenshot of the current screen. Returns file path and base64 image.",
            "parameters": {
                "type": "object",
                "properties": {
                    "region": {
                        "type": "object",
                        "description": "Optional region to capture: {x, y, w, h}. Omit for full screen.",
                        "properties": {
                            "x": {"type": "integer"},
                            "y": {"type": "integer"},
                            "w": {"type": "integer"},
                            "h": {"type": "integer"},
                        },
                    },
                },
            },
        },
    },
    "ocr_screen": {
        "fn": ocr_screen,
        "schema": {
            "name": "ocr_screen",
            "description": "Extract all text from the screen using Apple Vision OCR. Takes a screenshot and reads all visible text. Fast, offline, no model needed.",
            "parameters": {
                "type": "object",
                "properties": {
                    "image_path": {
                        "type": "string",
                        "description": "Path to an existing screenshot. If omitted, takes a fresh screenshot.",
                    },
                },
            },
        },
    },
    "ask_about_screen": {
        "fn": ask_about_screen,
        "schema": {
            "name": "ask_about_screen",
            "description": (
                "SCREENSHOT + DESCRIBE in a single call — captures the screen "
                "internally, then uses a vision model to answer the question. "
                "PICK THIS (not take_screenshot) whenever the user wants to know "
                "WHAT's on screen: 'what's on my screen', 'describe what you see', "
                "'take a screenshot and tell me what apps are open', 'what am I "
                "looking at', 'what's this error'. Handles code reading, app "
                "identification, UI explanation, error diagnosis. Do NOT call "
                "take_screenshot first — this does it internally."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "What to look for or answer about the screen. E.g. 'what app is open', 'read the error message', 'what language is this code'.",
                    },
                    "image_path": {
                        "type": "string",
                        "description": "Path to an existing screenshot. If omitted, takes a fresh screenshot.",
                    },
                },
                "required": ["query"],
            },
        },
    },
    "capture_full_page": {
        "fn": capture_full_page,
        "schema": {
            "name": "capture_full_page",
            "description": "Capture the FULL page by scrolling through the entire content and OCR-ing each viewport. Returns all text from the complete page, not just what's visible. Works with any scrollable app — browsers, PDFs, Word docs. Can target a specific app by name.",
            "parameters": {
                "type": "object",
                "properties": {
                    "max_scrolls": {
                        "type": "integer",
                        "description": "Maximum number of page-downs to attempt (default 20).",
                    },
                    "app": {
                        "type": "string",
                        "description": "App to activate before capturing (e.g. 'Safari', 'Preview', 'Google Chrome'). If omitted, uses the frontmost app.",
                    },
                },
            },
        },
    },
    "solve_screen_questions": {
        "fn": solve_screen_questions,
        "schema": {
            "name": "solve_screen_questions",
            "description": "Capture the screen (full page or current viewport), read all questions/problems, solve them, and save well-formatted answers to a .docx file. Can target a specific app. Use full_page=true for complete pages, full_page=false for just the current view.",
            "parameters": {
                "type": "object",
                "properties": {
                    "save_path": {
                        "type": "string",
                        "description": "Where to save the .docx. Defaults to ~/Documents/friday_files/Screen_Answers_<timestamp>.docx",
                    },
                    "app": {
                        "type": "string",
                        "description": "App to activate and capture from (e.g. 'Safari', 'Preview', 'Google Chrome'). If omitted, uses the frontmost app.",
                    },
                    "full_page": {
                        "type": "boolean",
                        "description": "If true (default), scrolls from top to capture the entire page. If false, only captures what's currently visible on screen.",
                    },
                },
            },
        },
    },
    "read_screen": {
        "fn": read_screen,
        "schema": {
            "name": "read_screen",
            "description": "Read all text content from the screen. General-purpose screen reader — reads job postings, articles, documents, forms, anything visible. Returns clean text with UI chrome stripped. Use this to understand what's on screen before taking action.",
            "parameters": {
                "type": "object",
                "properties": {
                    "app": {
                        "type": "string",
                        "description": "App to activate and read from (e.g. 'Safari', 'Preview'). If omitted, reads the frontmost app.",
                    },
                    "full_page": {
                        "type": "boolean",
                        "description": "If true, scrolls through entire page. If false (default), reads current viewport only.",
                    },
                },
            },
        },
    },
}
