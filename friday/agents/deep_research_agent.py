"""Deep Research Agent — multi-agent coordinator for complex tasks.

Not just research. This is the multi-agent pipeline for any task that needs
agents working together: research + write, read + improve, analyze + create.

Use cases:
- "deep research about X and save a paper on my desktop"
- "read my thesis, research its topics, improve it to research-paper grade"
- "I have an idea about X, research it, build a submission-ready document"
- Any task that needs: gathering → processing → producing a deliverable

Flow:
  1. Planner (1 LLM) — breaks task into a plan with typed steps
  2. Executor — runs steps in parallel where possible (read, search, fetch)
  3. Writer (parallel LLMs) — each section written simultaneously
  4. Assembler — merges into final document, saves to disk
"""

import asyncio
import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Optional, Callable

from friday.core.llm import cloud_chat, extract_text, extract_stream_content
from friday.core.types import AgentResponse, ToolResult
from friday.tools.web_tools import search_web, fetch_page


# ── Format utilities ─────────────────────────────────────────────────────────

SUPPORTED_FORMATS = {"docx", "md", "txt", "pdf"}


def detect_format(task: str) -> str:
    """Detect desired output format from the task description."""
    task_lower = task.lower()
    if any(kw in task_lower for kw in (".pdf", "pdf file", "as pdf", "in pdf")):
        return "pdf"
    if any(kw in task_lower for kw in (".md", "markdown", "as markdown", "in markdown")):
        return "md"
    if any(kw in task_lower for kw in (".txt", "text file", "as text", "plain text", "as txt")):
        return "txt"
    # Default: docx
    return "docx"


def save_document(
    save_path: Path,
    title: str,
    task: str,
    sections: list[dict],
    synth_text: str,
    all_sources: set,
    fmt: str = "docx",
):
    """Save document in the requested format."""
    if fmt == "docx":
        _save_docx(save_path, title, task, sections, synth_text, all_sources)
    elif fmt == "pdf":
        _save_pdf(save_path, title, sections, synth_text, all_sources)
    elif fmt == "md":
        _save_md(save_path, title, task, sections, synth_text, all_sources)
    elif fmt == "txt":
        _save_txt(save_path, title, sections, synth_text, all_sources)


def convert_file(source_path: str, target_format: str) -> Path:
    """Convert an existing file to another format."""
    src = Path(source_path).expanduser().resolve()
    if not src.exists():
        raise FileNotFoundError(f"File not found: {source_path}")

    target_format = target_format.lower().strip(".")
    if target_format not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {target_format}. Supported: {', '.join(SUPPORTED_FORMATS)}")

    dest = src.with_suffix(f".{target_format}")

    # Read source content based on its format
    suffix = src.suffix.lower()

    if suffix == ".md":
        content = src.read_text(encoding="utf-8")
        sections, title, synth, sources = _parse_md(content)
    elif suffix == ".txt":
        content = src.read_text(encoding="utf-8")
        sections = [{"title": "Content", "content": content}]
        title = src.stem.replace("_", " ")
        synth = ""
        sources = set()
    elif suffix == ".docx":
        sections, title, synth, sources = _parse_docx(src)
    else:
        # Try reading as text
        content = src.read_text(encoding="utf-8", errors="replace")
        sections = [{"title": "Content", "content": content}]
        title = src.stem.replace("_", " ")
        synth = ""
        sources = set()

    save_document(dest, title, "", sections, synth, sources, target_format)
    return dest


def _parse_md(content: str) -> tuple:
    """Parse markdown into sections, title, synth, sources."""
    lines = content.split("\n")
    title = "Document"
    sections = []
    current_title = None
    current_lines = []
    synth = ""
    sources = set()
    skip_sections = {"Abstract", "Table of Contents", "Conclusion", "References", "Sources"}

    for line in lines:
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
        elif line.startswith("## "):
            heading = line[3:].strip()
            if current_title and current_title not in skip_sections:
                sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})
            if heading == "Abstract":
                current_title = heading
                current_lines = []
            elif heading == "Conclusion":
                if current_title and current_title not in skip_sections:
                    pass  # already appended above
                current_title = heading
                current_lines = []
            elif heading in skip_sections:
                current_title = heading
                current_lines = []
            else:
                current_title = heading
                current_lines = []
        elif current_title:
            current_lines.append(line)

    if current_title and current_title not in skip_sections:
        sections.append({"title": current_title, "content": "\n".join(current_lines).strip()})

    # Build synth from abstract + conclusion
    abstract = conclusion = ""
    for line_block in content.split("## "):
        if line_block.startswith("Abstract"):
            abstract = line_block[len("Abstract"):].strip().split("\n---")[0].strip()
        if line_block.startswith("Conclusion"):
            conclusion = line_block[len("Conclusion"):].strip().split("\n---")[0].strip()
    if abstract:
        synth += f"## Abstract\n{abstract}\n\n"
    if conclusion:
        synth += f"## Conclusion\n{conclusion}"

    return sections, title, synth, sources


def _parse_docx(path: Path) -> tuple:
    """Parse docx into sections."""
    from docx import Document
    doc = Document(str(path))
    title = "Document"
    sections = []
    current_title = None
    current_paras = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split()[-1]) if para.style.name[-1].isdigit() else 0
            if level == 0:
                title = para.text
            elif level == 1:
                if current_title:
                    sections.append({"title": current_title, "content": "\n\n".join(current_paras)})
                current_title = para.text
                current_paras = []
            else:
                current_paras.append(f"### {para.text}")
        else:
            if para.text.strip():
                current_paras.append(para.text)

    if current_title:
        sections.append({"title": current_title, "content": "\n\n".join(current_paras)})

    return sections, title, "", set()


def _save_md(path: Path, title: str, task: str, sections: list, synth: str, sources: set):
    """Save as markdown."""
    now_str = datetime.now().strftime("%B %d, %Y")
    lines = [f"# {title}", "", f"*Compiled by FRIDAY — {now_str}*", "", "---", ""]

    if "## Abstract" in synth:
        abstract_part = synth.split("## Conclusion")[0] if "## Conclusion" in synth else synth
        lines.append(abstract_part.strip())
        lines.extend(["", "---", ""])

    lines.append("## Table of Contents")
    lines.append("")
    for i, s in enumerate(sections, 1):
        lines.append(f"{i}. {s['title']}")
    lines.extend(["", "---", ""])

    for s in sections:
        lines.extend([f"## {s['title']}", "", s["content"], "", "---", ""])

    if "## Conclusion" in synth:
        conclusion = synth.split("## Conclusion")[1].strip()
        lines.extend(["## Conclusion", "", conclusion, "", "---", ""])

    if sources:
        lines.append("## References")
        lines.append("")
        for i, src in enumerate(sorted(sources), 1):
            lines.append(f"{i}. {src}")

    path.write_text("\n".join(lines), encoding="utf-8")


def _save_txt(path: Path, title: str, sections: list, synth: str, sources: set):
    """Save as plain text."""
    lines = [title.upper(), "=" * len(title), ""]

    if "## Abstract" in synth:
        abstract = synth.split("## Abstract")[1].split("## Conclusion")[0].strip() if "## Conclusion" in synth else synth.split("## Abstract")[1].strip()
        lines.extend(["ABSTRACT", "-" * 8, "", abstract, "", ""])

    for s in sections:
        lines.extend([s["title"].upper(), "-" * len(s["title"]), "", s["content"], "", ""])

    if "## Conclusion" in synth:
        conclusion = synth.split("## Conclusion")[1].strip()
        lines.extend(["CONCLUSION", "-" * 10, "", conclusion, "", ""])

    if sources:
        lines.extend(["REFERENCES", "-" * 10, ""])
        for i, src in enumerate(sorted(sources), 1):
            lines.append(f"{i}. {src}")

    path.write_text("\n".join(lines), encoding="utf-8")


def _save_pdf(path: Path, title: str, sections: list, synth: str, sources: set):
    """Save as PDF. Uses WeasyPrint if available, falls back to basic text PDF."""
    # Build HTML content
    html_parts = [f"<h1>{title}</h1>"]

    if "## Abstract" in synth:
        abstract = synth.split("## Abstract")[1].split("## Conclusion")[0].strip() if "## Conclusion" in synth else synth.split("## Abstract")[1].strip()
        html_parts.append(f"<h2>Abstract</h2><p>{abstract}</p>")

    for s in sections:
        html_parts.append(f"<h2>{s['title']}</h2>")
        for para in s["content"].split("\n\n"):
            para = para.strip()
            if para.startswith("### "):
                html_parts.append(f"<h3>{para[4:]}</h3>")
            elif para:
                html_parts.append(f"<p>{para}</p>")

    if "## Conclusion" in synth:
        conclusion = synth.split("## Conclusion")[1].strip()
        html_parts.append(f"<h2>Conclusion</h2><p>{conclusion}</p>")

    if sources:
        html_parts.append("<h2>References</h2><ol>")
        for src in sorted(sources):
            html_parts.append(f"<li>{src}</li>")
        html_parts.append("</ol>")

    html = f"""<!DOCTYPE html>
<html><head><style>
body {{ font-family: Calibri, sans-serif; font-size: 11pt; line-height: 1.6; margin: 2cm; }}
h1 {{ text-align: center; font-size: 18pt; }}
h2 {{ font-size: 14pt; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
h3 {{ font-size: 12pt; }}
p {{ margin-bottom: 8px; }}
</style></head><body>{''.join(html_parts)}</body></html>"""

    try:
        from weasyprint import HTML as WeasyprintHTML
        WeasyprintHTML(string=html).write_pdf(str(path))
    except ImportError:
        # Fallback: save as HTML with .pdf extension note
        html_path = path.with_suffix(".html")
        html_path.write_text(html, encoding="utf-8")
        raise RuntimeError(
            f"WeasyPrint not installed for PDF generation. "
            f"Saved as HTML instead: {html_path}. Install with: uv add weasyprint"
        )


def _save_docx(path: Path, title: str, task: str, sections: list, synth: str, sources: set):
    """Save as .docx."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    style.paragraph_format.space_after = Pt(6)

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now_str = datetime.now().strftime("%B %d, %Y")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Compiled by FRIDAY — {now_str}")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    if "## Abstract" in synth:
        abstract_text = synth.split("## Abstract")[1]
        if "## Conclusion" in abstract_text:
            abstract_text = abstract_text.split("## Conclusion")[0]
        abstract_text = abstract_text.strip()
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(abstract_text)

    doc.add_heading("Table of Contents", level=1)
    for i, s in enumerate(sections, 1):
        toc_para = doc.add_paragraph(f"{i}. {s['title']}")
        toc_para.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    for s in sections:
        doc.add_heading(s["title"], level=1)
        for para_text in s["content"].split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("### "):
                doc.add_heading(para_text[4:], level=2)
            elif para_text.startswith("## "):
                doc.add_heading(para_text[3:], level=2)
            elif para_text.startswith("**") and para_text.endswith("**"):
                doc.add_heading(para_text.strip("*"), level=2)
            else:
                doc.add_paragraph(para_text)

    if "## Conclusion" in synth:
        conclusion = synth.split("## Conclusion")[1].strip()
        doc.add_heading("Conclusion", level=1)
        for para_text in conclusion.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                doc.add_paragraph(para_text)

    if sources:
        doc.add_heading("References", level=1)
        for i, src in enumerate(sorted(sources), 1):
            doc.add_paragraph(f"{i}. {src}")

    doc.save(str(path))


# ── Prompts ──────────────────────────────────────────────────────────────────

PLANNER_PROMPT = """You are a task planner for a multi-agent AI system. Given a complex task, break it into a structured plan.

You have these capabilities:
- SEARCH: Web search for information (query)
- FETCH: Fetch a specific URL for full content
- READ_FILE: Read a local file (path)
- WRITE: Write a section of a document (topic + data)

Break the task into:
1. A document title
2. An ordered list of steps, each with a type and parameters
3. Group steps that can run in parallel

Respond ONLY in this JSON format:
{
  "title": "Document title",
  "output_format": "md",
  "steps": [
    {
      "phase": 1,
      "type": "READ_FILE",
      "params": {"path": "/path/to/file"},
      "description": "Read the existing thesis"
    },
    {
      "phase": 2,
      "type": "SEARCH",
      "params": {"queries": ["query 1", "query 2", "query 3"]},
      "description": "Research topic A"
    },
    {
      "phase": 2,
      "type": "SEARCH",
      "params": {"queries": ["query 4", "query 5"]},
      "description": "Research topic B"
    },
    {
      "phase": 3,
      "type": "WRITE",
      "params": {"section_title": "Section Name"},
      "description": "Write section based on research"
    }
  ],
  "sections": ["Introduction", "Section A", "Section B", "Conclusion"]
}

Rules:
- Steps in the SAME phase run in PARALLEL (fast)
- Steps in LATER phases wait for earlier ones (dependencies)
- Always include 4-8 SEARCH steps with 2-3 queries each for thorough coverage
- If the task mentions reading a file, include a READ_FILE step in phase 1
- SEARCH queries should be specific and diverse — cover different angles
- sections list = the final document structure (in order)
- Do NOT include "Introduction", "Conclusion", "Summary", or "Abstract" in the sections list — those are generated automatically by the synthesis step"""

SECTION_WRITER_PROMPT = """You are an expert writer. Write a detailed, well-structured section for a document.

Rules:
- Write 400-600 words for this section
- Be specific — include facts, numbers, names, dates, technical details
- Cite sources inline using their full URL: [Source: https://example.com/page]. The research data includes [Source: URL] tags — use those exact URLs, not just the domain or "pdf"
- Professional but accessible tone — clear, not stuffy
- If writing for academic submission: use formal language, structured arguments, evidence-based claims
- If improving existing text: preserve the author's voice while strengthening arguments and adding depth
- Do NOT use placeholder text or filler
- Do NOT repeat information from other sections
- Do NOT write a conclusion or summary — that is handled separately
- Do NOT start with the section title as a heading — just write the content directly"""

SYNTHESIS_PROMPT = """You are an expert writer. Given all the sections below, write:
1. An executive summary / abstract (150-250 words) covering the key findings and arguments
2. A conclusion (200-300 words) synthesizing insights, implications, and recommendations

Document title: {title}
Original task: {task}

Sections written so far:
{sections}

Respond in this format:
## Abstract
[summary text]

## Conclusion
[conclusion text]"""

IMPROVE_PROMPT = """You are an expert editor and researcher. You have been given an existing document and new research data.

Your job: improve the document to research-paper grade. Specifically:
- Strengthen arguments with evidence from the research data
- Add citations and references
- Improve structure and flow
- Add missing perspectives or counterarguments
- Make the language more precise and academic
- Keep the original author's core ideas and voice

Original document:
{original}

New research data:
{research}

Write the improved version of this section. Be thorough."""


class DeepResearchAgent:
    name = "deep_research_agent"

    async def run(
        self,
        task: str,
        context: str = "",
        on_tool_call: Optional[Callable] = None,
        on_chunk: Optional[Callable] = None,
    ) -> AgentResponse:
        t0 = time.time()
        tools_called = []

        def _status(msg):
            if on_chunk:
                on_chunk(f"\n  ◈ {msg}")

        # ── Step 1: Plan (1 LLM call) ──
        _status("Planning task structure...")

        plan_messages = [
            {"role": "system", "content": PLANNER_PROMPT},
            {"role": "user", "content": task},
        ]
        plan_response = cloud_chat(messages=plan_messages, max_tokens=1200)
        plan_text = extract_text(plan_response)

        try:
            json_text = plan_text
            if "```" in json_text:
                json_text = json_text.split("```")[1]
                if json_text.startswith("json"):
                    json_text = json_text[4:]
            plan = json.loads(json_text.strip())
        except (json.JSONDecodeError, IndexError):
            return AgentResponse(
                agent_name=self.name, success=False,
                result=f"Couldn't parse plan. Raw:\n{plan_text[:500]}",
                tools_called=[], duration_ms=int((time.time() - t0) * 1000),
            )

        title = plan.get("title", "Research Document")
        steps = plan.get("steps", [])
        section_names = plan.get("sections", [])

        if not steps:
            return AgentResponse(
                agent_name=self.name, success=False,
                result="Planner returned no steps.",
                tools_called=[], duration_ms=int((time.time() - t0) * 1000),
            )

        _status(f"Plan: {len(steps)} steps across {max(s.get('phase', 1) for s in steps)} phases")
        for s in steps:
            _status(f"  Phase {s.get('phase', '?')}: [{s['type']}] {s.get('description', '')}")

        # ── Step 2: Execute steps phase by phase ──
        # Group by phase, run each phase in parallel
        phases = {}
        for step in steps:
            phase = step.get("phase", 1)
            phases.setdefault(phase, []).append(step)

        # Accumulated context from all phases
        gathered_data = {}  # description -> result text
        original_content = None  # if we read a file

        for phase_num in sorted(phases.keys()):
            phase_steps = phases[phase_num]
            _status(f"Phase {phase_num}: running {len(phase_steps)} steps in parallel...")

            async def _execute_step(step: dict) -> tuple:
                step_type = step["type"]
                params = step.get("params", {})
                desc = step.get("description", step_type)

                if step_type == "SEARCH":
                    return desc, await self._execute_search(
                        params.get("queries", []),
                        tools_called,
                    )

                elif step_type == "FETCH":
                    url = params.get("url", "")
                    if url:
                        tools_called.append("fetch_page")
                        result = await fetch_page(url)
                        if isinstance(result, ToolResult) and result.success:
                            data = result.data
                            content = data.get("content", data.get("text", "")) if isinstance(data, dict) else str(data)
                            return desc, content[:5000]
                    return desc, ""

                elif step_type == "READ_FILE":
                    path = params.get("path", "")
                    if path:
                        tools_called.append("read_file")
                        try:
                            from friday.tools.file_tools import read_file
                            result = await read_file(path)
                            if isinstance(result, ToolResult) and result.success:
                                content = result.data if isinstance(result.data, str) else result.data.get("content", str(result.data))
                                return desc, content[:10000]
                        except Exception as e:
                            return desc, f"Failed to read file: {e}"
                    return desc, ""

                elif step_type == "WRITE":
                    # Write steps are handled in the writing phase
                    return desc, ""

                return desc, ""

            results = await asyncio.gather(
                *[_execute_step(s) for s in phase_steps],
                return_exceptions=True,
            )

            for r in results:
                if isinstance(r, Exception):
                    continue
                desc, data = r
                if data:
                    gathered_data[desc] = data
                    # Track if we read an original file
                    if any(s["type"] == "READ_FILE" and s.get("description") == desc for s in phase_steps):
                        original_content = data

        _status(f"Data gathered: {len(gathered_data)} sources, {sum(len(v) for v in gathered_data.values())} chars")

        # ── Step 3: Write sections in parallel ──
        if not section_names:
            section_names = [s.get("description", f"Section {i}") for i, s in enumerate(steps) if s["type"] == "SEARCH"]
            if not section_names:
                section_names = list(gathered_data.keys())

        # Filter out sections that synthesis handles (prevents duplicate conclusion/intro)
        _skip = {"introduction", "conclusion", "summary", "abstract",
                 "conclusion and summary", "references", "table of contents"}
        section_names = [s for s in section_names if s.lower().strip() not in _skip]

        _status(f"Writing {len(section_names)} sections...")

        # Build research context for writers — all gathered data merged
        all_research = "\n\n---\n\n".join(
            f"### {desc}\n{data[:4000]}" for desc, data in gathered_data.items()
        )

        # Decide if this is an improvement task or a fresh write
        is_improvement = original_content is not None

        async def _write_section(section_title: str, section_idx: int) -> dict:
            if is_improvement:
                # Extract the relevant part of the original if possible
                messages = [
                    {"role": "system", "content": IMPROVE_PROMPT.format(
                        original=original_content[:3000],
                        research=all_research[:6000],
                    )},
                    {"role": "user", "content": (
                        f"Write/improve this section: {section_title}\n\n"
                        f"This is section {section_idx + 1} of {len(section_names)}. "
                        f"Focus on the aspects relevant to '{section_title}'."
                    )},
                ]
            else:
                # Partition research data — give each section the most relevant chunks
                # Simple: give all data but tell it which section to focus on
                messages = [
                    {"role": "system", "content": SECTION_WRITER_PROMPT},
                    {"role": "user", "content": (
                        f"Section title: {section_title}\n"
                        f"This is section {section_idx + 1} of {len(section_names)} "
                        f"in a document titled \"{title}\".\n"
                        f"Original task: {task}\n\n"
                        f"Research data:\n{all_research[:6000]}\n\n"
                        f"Write a detailed, well-structured section focused specifically on "
                        f"'{section_title}'. Do not cover topics from other sections."
                    )},
                ]

            response = cloud_chat(messages=messages, max_tokens=1200)
            text = extract_text(response)
            # Strip leading heading that duplicates the section title
            # (LLMs often echo "## Section Title" at the start)
            for prefix in (f"## {section_title}", f"# {section_title}", f"**{section_title}**"):
                if text.strip().startswith(prefix):
                    text = text.strip()[len(prefix):].strip()
                    break
            return {"title": section_title, "content": text}

        # Stagger LLM calls to avoid Groq rate limits (6K TPM)
        # Run in batches of 3 with a small delay between batches
        section_results = []
        batch_size = 3
        for batch_start in range(0, len(section_names), batch_size):
            batch = section_names[batch_start:batch_start + batch_size]
            batch_results = await asyncio.gather(
                *[_write_section(name, batch_start + i) for i, name in enumerate(batch)],
                return_exceptions=True,
            )
            section_results.extend(batch_results)
            if batch_start + batch_size < len(section_names):
                await asyncio.sleep(1)  # Brief pause between batches

        sections = []
        for r in section_results:
            if isinstance(r, Exception):
                continue
            if isinstance(r, dict):
                sections.append(r)

        _status(f"Sections written: {len(sections)}/{len(section_names)}")

        if not sections:
            return AgentResponse(
                agent_name=self.name, success=False,
                result="All section writers failed.",
                tools_called=tools_called, duration_ms=int((time.time() - t0) * 1000),
            )

        # ── Step 4: Abstract + Conclusion (1 LLM call) ──
        _status("Writing abstract and conclusion...")

        sections_text = "\n\n".join(f"## {s['title']}\n{s['content']}" for s in sections)

        synth_messages = [
            {"role": "system", "content": SYNTHESIS_PROMPT.format(
                title=title, task=task, sections=sections_text[:8000],
            )},
            {"role": "user", "content": "Write the abstract and conclusion."},
        ]
        synth_response = cloud_chat(messages=synth_messages, max_tokens=1000)
        synth_text = extract_text(synth_response)

        # ── Step 5: Assemble document ──
        _status("Assembling final document...")

        # Collect sources from research data AND written sections
        all_sources = set()
        for data in gathered_data.values():
            urls = re.findall(r'https?://[^\s\])<>"]+', data)
            all_sources.update(urls[:10])
        # Also extract [Source: domain] citations from written sections
        for s in sections:
            source_refs = re.findall(r'\[Source:\s*([^\]]+)\]', s.get("content", ""))
            for ref in source_refs:
                ref = ref.strip()
                if not ref.startswith("http"):
                    ref = f"https://{ref}"
                all_sources.add(ref)

        now_str = datetime.now().strftime("%B %d, %Y")
        doc_lines = [
            f"# {title}",
            "",
            f"*Compiled by FRIDAY — {now_str}*",
            f"*Task: {task}*",
            "",
            "---",
            "",
        ]

        # Abstract
        if "## Abstract" in synth_text:
            abstract_part = synth_text.split("## Conclusion")[0] if "## Conclusion" in synth_text else synth_text
            doc_lines.append(abstract_part.strip())
            doc_lines.extend(["", "---", ""])

        # Table of contents
        doc_lines.append("## Table of Contents")
        doc_lines.append("")
        for i, s in enumerate(sections, 1):
            doc_lines.append(f"{i}. {s['title']}")
        doc_lines.extend(["", "---", ""])

        # Sections
        for s in sections:
            doc_lines.append(f"## {s['title']}")
            doc_lines.append("")
            doc_lines.append(s["content"])
            doc_lines.extend(["", "---", ""])

        # Conclusion
        if "## Conclusion" in synth_text:
            conclusion = synth_text.split("## Conclusion")[1].strip()
            doc_lines.append("## Conclusion")
            doc_lines.append("")
            doc_lines.append(conclusion)
            doc_lines.extend(["", "---", ""])

        # References
        if all_sources:
            doc_lines.append("## References")
            doc_lines.append("")
            for i, src in enumerate(sorted(all_sources), 1):
                doc_lines.append(f"{i}. {src}")
            doc_lines.append("")

        document = "\n".join(doc_lines)

        # ── Step 6: Save to disk ──
        fmt = detect_format(task)
        save_path = self._determine_save_path(task, title, fmt)
        save_document(save_path, title, task, sections, synth_text, all_sources, fmt)

        duration = int((time.time() - t0) * 1000)
        elapsed_s = duration / 1000

        summary = (
            f"Done. {len(sections)} sections, {len(all_sources)} sources, "
            f"{len(document)} chars. Saved to {save_path}. "
            f"({elapsed_s:.0f}s, {len(tools_called)} tool calls)"
        )

        _status(summary)

        # Stream preview
        if on_chunk:
            preview = document[:2000]
            on_chunk(f"\n\n{preview}")
            if len(document) > 2000:
                on_chunk(f"\n\n... (full document at {save_path})")

        return AgentResponse(
            agent_name=self.name,
            success=True,
            result=summary,
            data={"path": str(save_path), "document_length": len(document)},
            tools_called=tools_called,
            duration_ms=duration,
        )

    async def _execute_search(self, queries: list, tools_called: list) -> str:
        """Run multiple search queries in parallel, merge results."""
        if not queries:
            return ""

        search_tasks = [search_web(q, num_results=5) for q in queries]
        search_results = await asyncio.gather(*search_tasks, return_exceptions=True)

        texts = []
        urls = set()

        for result in search_results:
            tools_called.append("search_web")
            if isinstance(result, Exception):
                continue
            if isinstance(result, ToolResult) and result.success and result.data:
                data = result.data
                if isinstance(data, dict):
                    if data.get("answer"):
                        texts.append(data["answer"])
                    for r in data.get("results", []):
                        url = r.get("url", "")
                        title = r.get("title", "")
                        content = r.get("content", "")
                        # Tag every snippet with its source URL so writers can cite properly
                        source_tag = f" [Source: {url}]" if url else ""
                        texts.append(f"{title}: {content}{source_tag}")
                        if url:
                            urls.add(url)
                elif isinstance(data, str):
                    texts.append(data)

        # Fetch top 2 pages for deeper content — tag with source URL
        fetch_urls = list(urls)[:2]
        if fetch_urls:
            fetch_tasks = [fetch_page(url) for url in fetch_urls]
            fetch_results = await asyncio.gather(*fetch_tasks, return_exceptions=True)
            for url, result in zip(fetch_urls, fetch_results):
                tools_called.append("fetch_page")
                if isinstance(result, Exception):
                    continue
                if isinstance(result, ToolResult) and result.success and result.data:
                    page_data = result.data
                    content = page_data.get("content", page_data.get("text", "")) if isinstance(page_data, dict) else str(page_data)
                    if content:
                        texts.append(f"[Source: {url}]\n{content[:3000]}")

        return "\n\n".join(texts)[:8000]

    @staticmethod
    def _determine_save_path(task: str, title: str, fmt: str = "docx") -> Path:
        """Figure out where to save based on the task wording."""
        task_lower = task.lower()

        # Check if user specified a location
        if "desktop" in task_lower:
            save_dir = Path.home() / "Desktop"
        elif "download" in task_lower:
            save_dir = Path.home() / "Downloads"
        else:
            # Default: ~/Documents/friday_files
            save_dir = Path.home() / "Documents" / "friday_files"
            save_dir.mkdir(parents=True, exist_ok=True)

        # Clean title for filename
        safe_title = "".join(c if c.isalnum() or c in " -_" else "" for c in title)
        safe_title = safe_title.strip().replace(" ", "_")[:60]
        if not safe_title:
            safe_title = "research_paper"

        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.{fmt}"
        return save_dir / filename

